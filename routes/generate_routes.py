from flask import Blueprint, request, jsonify, current_app, send_file
from flask_jwt_extended import jwt_required
import google.generativeai as genai
from docx import Document
from ultis.pdf_ultis import read_pdf_text  
from controllers.generate_controller import generate_lesson_content, parse_markdown_to_json
import io

generate_bp = Blueprint('generate', __name__)

# --- Model cấu hình ---
def get_model():
    api_key = current_app.config.get('GEMINI_API_KEY')
    if not api_key:
        raise ValueError("GEMINI_API_KEY chưa được cấu hình")
    genai.configure(api_key=api_key)
    return genai.GenerativeModel("gemini-1.5-flash-latest")

# --- API tạo dàn ý ---
@generate_bp.route('/outline', methods=['POST'])
def generate_outline():
    text = request.json.get("text")
    if not text:
        return jsonify({"error": "Thiếu dữ liệu"}), 400
    try:
        model = get_model()
        response = model.generate_content(f"Hãy tạo dàn ý cho bài giảng: {text}")
        return jsonify({"input": text, "result": response.text})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# --- API tạo giáo án ---
# controllers/generate_controller.py
# Bạn dùng Gemini API ở đây
@generate_bp.route("/api/generate/lesson", methods=["POST"])
def generate_lesson():
    data = request.json
    user_input = data.get("topic", "")
    pdf_path = "C:/Users/minhs/OneDrive/Videos/SGK KHTN 8 KNTT.pdf"

    result = generate_lesson_from_pdf(pdf_path, user_input)
    return jsonify({"result": result})


# --- API tạo slide ---
@generate_bp.route('/slides', methods=['POST'])
def generate_slides():
    text = request.json.get("text")
    if not text:
        return jsonify({"error": "Thiếu dữ liệu"}), 400
    try:
        model = get_model()
        response = model.generate_content(f"Hãy tạo slide trình bày cho bài giảng sau: {text}")
        return jsonify({"input": text, "result": response.text})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# --- API tóm tắt tài liệu ---
@generate_bp.route('/summarize', methods=['POST'])
def generate_summarize():
    text = request.json.get("text")
    if not text:
        return jsonify({"error": "Thiếu dữ liệu"}), 400
    try:
        model = get_model()
        response = model.generate_content(f"Hãy tóm tắt nội dung sau: {text}")
        return jsonify({"input": text, "result": response.text})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# --- API chat từng câu (giống ChatGPT) ---
@generate_bp.route('/chat/message', methods=['POST'])
@jwt_required()
def chat_message():
    data = request.get_json()
    message = data.get("message", "")
    if not message:
        return jsonify({"error": "Thiếu message"}), 400
    try:
        model = get_model()
        response = model.generate_content(message)
        return jsonify({"reply": response.text})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# --- API chat theo luồng: tạo giáo án từ hội thoại ---
@generate_bp.route('/chat/flow', methods=['POST'])
def chat_flow():
    data = request.json
    user_message = data.get("message", "")
    state = data.get("state", "ask_topic")

    # Mặc định trả về các trường rỗng nếu chưa có
    outline = None
    lesson = None
    slides = None

    try:
        model = get_model()

        if state == "ask_topic":
            reply = "Chào thầy/cô! Hôm nay bạn muốn dạy chủ đề gì?"
            next_state = "wait_topic"

        elif state == "wait_topic":
            response = model.generate_content(f"Hãy tạo dàn ý bài giảng cho chủ đề: {user_message}")
            outline = response.text
            reply = f"Đây là dàn ý gợi ý:\n{outline}\n\nBạn thấy ổn chưa? (Nhập 'Tiếp tục' để tạo giáo án hoặc chỉnh sửa nếu cần)"
            next_state = "confirm_outline"

        elif state == "confirm_outline":
            if "tiếp tục" in user_message.lower():
                outline = data.get("outline", "")
                response = model.generate_content(f"Hãy viết giáo án chi tiết từ dàn ý sau:\n{outline}")
                lesson = response.text
                reply = f"Đây là giáo án chi tiết:\n{lesson}\n\nBạn muốn tạo slide không?"
                next_state = "confirm_lesson"
            else:
                reply = "Bạn muốn chỉnh sửa dàn ý thế nào? Hãy nhập nội dung sửa."
                next_state = "edit_outline"

        elif state == "edit_outline":
            outline = user_message
            response = model.generate_content(f"Hãy viết giáo án chi tiết từ dàn ý sau:\n{outline}")
            lesson = response.text
            reply = f"Đây là giáo án mới dựa trên dàn ý đã sửa:\n{lesson}\n\nBạn muốn tạo slide không?"
            next_state = "confirm_lesson"

        elif state == "confirm_lesson":
            if "có" in user_message.lower() or "vâng" in user_message.lower():
                lesson = data.get("lesson", "")
                response = model.generate_content(f"Hãy tạo slide trình bày cho bài giảng sau:\n{lesson}")
                slides = response.text
                reply = f"Dưới đây là slide trình bày:\n{slides}\n\nHoàn tất quá trình rồi nhé!"
                next_state = "done"
            else:
                reply = "Đã dừng tại bước giáo án. Bạn có thể yêu cầu tạo slide sau."
                next_state = "done"

        else:
            reply = "Kết thúc hội thoại. Nếu muốn bắt đầu lại, hãy nhập 'bắt đầu'."
            next_state = "ask_topic"

        return jsonify({
            "reply": reply,
            "next_state": next_state,
            "outline": outline,
            "lesson": lesson,
            "slides": slides
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# --- API xuất file Word từ nội dung ---
@generate_bp.route('/export-word', methods=['POST'])
def export_word():
    data = request.get_json()
    content = data.get('content')
    if not content:
        return jsonify({'error': 'Thiếu nội dung'}), 400

    doc = Document()

    # Tách các đoạn theo 2 dòng xuống hàng (giả sử mỗi mục lớn cách nhau 2 dòng)
    sections = [s.strip() for s in content.split('\n\n') if s.strip()]
    for section in sections:
        lines = section.split('\n')
        # Nếu đoạn là tiêu đề lớn (in hoa, ngắn), dùng heading
        if len(lines) == 1 and (lines[0].isupper() or len(lines[0]) < 40):
            doc.add_heading(lines[0], level=1)
            continue

        # Nếu đoạn là bảng (dòng có nhiều dấu "|"), tạo bảng
        if all('|' in line for line in lines):
            rows = [l.split('|') for l in lines]
            table = doc.add_table(rows=1, cols=len(rows[0]))
            table.style = 'Table Grid'
            # Header
            for i, cell in enumerate(rows[0]):
                table.cell(0, i).text = cell.strip()
            # Body
            for row in rows[1:]:
                cells = table.add_row().cells
                for i, cell in enumerate(row):
                    cells[i].text = cell.strip()
            continue

        # Nếu là danh sách (bắt đầu bằng -, *, •, số), tạo bullet/numbered list
        if all(line.strip().startswith(("-", "*", "•", "1.", "2.", "3.", "4.", "5.")) for line in lines):
            for line in lines:
                doc.add_paragraph(line.strip(), style='List Bullet')
            continue

        # Nếu là mục lớn (bắt đầu bằng số/chữ cái và dấu chấm), in đậm
        if lines[0][:2].replace('.', '').isnumeric() or lines[0][:2].isalpha():
            p = doc.add_paragraph()
            run = p.add_run(lines[0])
            run.bold = True
            for line in lines[1:]:
                doc.add_paragraph(line)
            continue

        # Mặc định: đoạn văn thường
        for line in lines:
            doc.add_paragraph(line)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name="giao_an.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )